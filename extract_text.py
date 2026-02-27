import os
import sys

def extract_pdf(pdf_path, output_txt):
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        with open(output_txt, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Successfully extracted {pdf_path} to {output_txt}")
    except Exception as e:
        print(f"Error extracting PDF: {e}")

def extract_doc(doc_path, output_txt):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        abs_doc_path = os.path.abspath(doc_path)
        print(f"Opening {abs_doc_path}")
        doc = word.Documents.Open(abs_doc_path)
        text = doc.Content.Text
        
        with open(output_txt, "w", encoding="utf-8") as f:
            f.write(text)
            
        doc.Close()
        word.Quit()
        print(f"Successfully extracted {doc_path} to {output_txt}")
    except Exception as e:
        print(f"Error extracting DOC: {e}")
        try:
            word.Quit()
        except:
            pass
        # Fallback to python-docx if it's actually a docx masquerading as doc
        try:
            import docx
            doc = docx.Document(doc_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join([cell.text for cell in row.cells]) + "\n"
            with open(output_txt, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"Successfully extracted {doc_path} (via docx) to {output_txt}")
        except Exception as e2:
            print(f"Fallback docx error: {e2}")


if __name__ == "__main__":
    pdf_file = "信息系统项目管理课程大纲.pdf"
    doc_file = "教学日历_信息系统项目管理_杨辉_信管2301.doc"
    
    extract_pdf(pdf_file, "syllabus.txt")
    extract_doc(doc_file, "calendar.txt")
